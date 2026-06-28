# =========== Copyright 2026 @ WANXIANG. All Rights Reserved. ===========
"""上传文档 → 纯文本提取 + LLM 提炼成模拟素材(material)。

设计:
- 文本类(pdf/docx/xlsx/txt/md/csv)由轻量纯 Python 库提取文字;图片走视觉
  模型读图。两者最终都经 ``distill_material`` 由 LLM 提炼成简洁素材。
- 重依赖(pypdf/python-docx/openpyxl)只在本模块按需 import,**不污染核心
  模拟链路**——运行时镜像核心保持精简(见 Dockerfile 注释)。
"""
from __future__ import annotations

import base64
import io

from wanxiang.simulation.decision import ModelCall

# 提取文本的硬上限(字符),防超大文档拖垮 LLM / 内存。
MAX_EXTRACT_CHARS = 20_000

TEXT_EXTS = {"txt", "md", "markdown", "csv", "tsv", "log", "json"}
IMAGE_EXTS = {"png", "jpg", "jpeg", "webp", "gif", "bmp"}
DOC_EXTS = {"pdf", "docx", "xlsx"}
SUPPORTED_EXTS = TEXT_EXTS | IMAGE_EXTS | DOC_EXTS


class DocParseError(Exception):
    """文档解析失败(格式不支持 / 提不出文字 / 损坏)。"""


def ext_of(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def kind_of(filename: str) -> str:
    """返回 'text' | 'image' | 'unsupported'(图片单独走视觉模型)。"""
    e = ext_of(filename)
    if e in IMAGE_EXTS:
        return "image"
    if e in TEXT_EXTS or e in DOC_EXTS:
        return "text"
    return "unsupported"


def _truncate(s: str) -> str:
    s = s.strip()
    if len(s) > MAX_EXTRACT_CHARS:
        return s[:MAX_EXTRACT_CHARS] + "\n…(内容过长已截断)"
    return s


def extract_text(filename: str, data: bytes) -> str:
    """从文本类文档提取纯文本。失败抛 DocParseError。"""
    e = ext_of(filename)

    if e in TEXT_EXTS:
        for enc in ("utf-8", "gbk", "latin-1"):
            try:
                return _truncate(data.decode(enc))
            except UnicodeDecodeError:
                continue
        raise DocParseError("无法识别文本编码")

    if e == "pdf":
        try:
            from pypdf import PdfReader
        except ImportError as ex:  # pragma: no cover
            raise DocParseError(f"PDF 解析库未安装: {ex}")
        try:
            reader = PdfReader(io.BytesIO(data))
            parts = [(pg.extract_text() or "") for pg in reader.pages]
        except Exception as ex:
            raise DocParseError(f"PDF 解析失败: {ex}")
        text = "\n".join(p for p in parts if p.strip())
        if not text.strip():
            raise DocParseError(
                "该 PDF 没有可提取的文字(可能是扫描版),可改传图片")
        return _truncate(text)

    if e == "docx":
        try:
            import docx  # python-docx
        except ImportError as ex:  # pragma: no cover
            raise DocParseError(f"Word 解析库未安装: {ex}")
        try:
            doc = docx.Document(io.BytesIO(data))
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            # 表格文字
            for tbl in doc.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paras.append(" | ".join(cells))
        except Exception as ex:
            raise DocParseError(f"Word 解析失败: {ex}")
        text = "\n".join(paras)
        if not text.strip():
            raise DocParseError("该 Word 文档没有可提取的文字")
        return _truncate(text)

    if e == "xlsx":
        try:
            from openpyxl import load_workbook
        except ImportError as ex:  # pragma: no cover
            raise DocParseError(f"Excel 解析库未安装: {ex}")
        try:
            wb = load_workbook(io.BytesIO(data), read_only=True,
                               data_only=True)
            lines: list[str] = []
            for ws in wb.worksheets:
                lines.append(f"# {ws.title}")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        lines.append(" | ".join(cells))
        except Exception as ex:
            raise DocParseError(f"Excel 解析失败: {ex}")
        text = "\n".join(lines)
        if not text.strip():
            raise DocParseError("该 Excel 没有可提取的内容")
        return _truncate(text)

    raise DocParseError(f"不支持的格式: .{e}")


_DISTILL_SYS = (
    "你是市场调研助手。请从用户提供的资料中提炼用于「人群购买/反应预测」的"
    "产品或投放素材:核心卖点、产品定位、价格/规格、目标人群线索、关键文案。"
    "去掉无关排版、页眉页脚、目录。用简洁中文输出,300 字以内,不要加解释性"
    "前缀,直接给素材正文。"
)


async def distill_material(text: str, model_call: ModelCall) -> str:
    """文本类:把提取的全文交给 LLM 提炼成简洁素材。"""
    messages = [
        {"role": "system", "content": _DISTILL_SYS},
        {"role": "user", "content": text},
    ]
    out = (await model_call(messages)).strip()
    return out or text[:500]


async def distill_image_material(data: bytes, filename: str,
                                 model_call: ModelCall) -> str:
    """图片:走视觉模型(image_url block)读图并提炼素材。

    需要工作区配置视觉模型;若模型不支持图片,model_call 会抛错,由调用方
    转成明确提示。
    """
    e = ext_of(filename)
    mime = "image/jpeg" if e in ("jpg", "jpeg") else f"image/{e or 'png'}"
    b64 = base64.b64encode(data).decode("ascii")
    messages = [
        {"role": "system", "content": _DISTILL_SYS},
        {"role": "user", "content": [
            {"type": "text",
             "text": "这是产品/投放资料图片,请提炼成预测素材。"},
            {"type": "image_url",
             "image_url": {"url": f"data:{mime};base64,{b64}"}},
        ]},
    ]
    out = (await model_call(messages)).strip()
    if not out:
        raise DocParseError("视觉模型未能从图片提炼出内容")
    return out
